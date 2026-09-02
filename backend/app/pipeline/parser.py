"""文档格式解析层：PDF / Word / Markdown / HTML / TXT 统一解析为结构化块。

保留标题层级、表格、列表结构。可选依赖按需延迟导入，缺失时降级。
"""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Optional, Union

from app.utils.logging import get_logger

logger = get_logger(__name__)

# 仅匹配“章/节/部分/篇/编”等标题，不匹配“条/款”（条款属于正文，交由切片器处理）
HEADING_RE = re.compile(r"^(第[一二三四五六七八九十百0-9]+[章节部分篇编]|[一二三四五六七八九十]+、|[0-9]+[.、．]\s*[^0-9])")


@dataclass
class Block:
    type: str  # heading | paragraph | list_item | table
    level: int  # 标题层级 1-6
    text: str
    page: Optional[int] = None
    extra: dict[str, Any] = field(default_factory=dict)


@dataclass
class ParsedDocument:
    title: str
    blocks: list[Block]
    raw_text: str = ""
    page_count: int = 0
    meta: dict[str, Any] = field(default_factory=dict)

    @property
    def text(self) -> str:
        if self.raw_text:
            return self.raw_text
        return "\n".join(b.text for b in self.blocks if b.text)


class DocumentParser:
    """按文件类型分派解析器。"""

    def parse(self, path: Union[str, Path]) -> ParsedDocument:
        path = Path(path)
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            return self._parse_pdf(path)
        if suffix == ".docx":
            return self._parse_docx(path)
        if suffix in (".md", ".markdown"):
            return self._parse_markdown(path)
        if suffix in (".html", ".htm"):
            return self._parse_html(path)
        if suffix in (".txt", ""):
            return self._parse_text(path)
        raise ValueError(f"不支持的文件类型: {suffix}")

    # ---------- PDF ----------
    def _parse_pdf(self, path: Path) -> ParsedDocument:
        text_pages: list[str] = []
        try:
            import pdfplumber  # 延迟导入

            with pdfplumber.open(str(path)) as pdf:
                for page in pdf.pages:
                    text_pages.append(page.extract_text() or "")
        except ImportError:
            text_pages = self._parse_pdf_pypdf(path)
        except Exception as exc:  # noqa: BLE001
            logger.warning("pdfplumber 解析失败(%s)，回退 pypdf", exc)
            text_pages = self._parse_pdf_pypdf(path)

        blocks = self._blocks_from_pages(text_pages)
        title = path.stem
        return ParsedDocument(title=title, blocks=blocks, raw_text="\n".join(text_pages), page_count=len(text_pages))

    def _parse_pdf_pypdf(self, path: Path) -> list[str]:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        return [(page.extract_text() or "") for page in reader.pages]

    # ---------- Word ----------
    def _parse_docx(self, path: Path) -> ParsedDocument:
        import docx  # 延迟导入

        doc = docx.Document(str(path))
        blocks: list[Block] = []
        # 文档内段落遍历（含表格）
        from docx.document import Document as _Doc
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        title = path.stem
        body = doc.element.body
        for child in body.iterchildren():
            if child.tag.endswith("}p"):
                para = Paragraph(child, doc)
                blocks.append(self._paragraph_to_block(para))
            elif child.tag.endswith("}tbl"):
                table = Table(child, doc)
                blocks.append(self._table_to_block(table))
        raw = "\n".join(b.text for b in blocks if b.text)
        return ParsedDocument(title=title, blocks=blocks, raw_text=raw, meta={"format": "docx"})

    def _paragraph_to_block(self, para: Any) -> Block:
        text = para.text.strip()
        style_name = (para.style.name if para.style and para.style.name else "") or ""
        level = 0
        if "Heading" in style_name or "标题" in style_name:
            digits = re.findall(r"\d+", style_name)
            level = int(digits[0]) if digits else 1
            return Block(type="heading", level=level, text=text)
        if text.startswith(("•", "- ", "·", "1.", "2.", "3.", "（", "(")):
            return Block(type="list_item", level=1, text=text)
        return Block(type="paragraph", level=0, text=text)

    def _table_to_block(self, table: Any) -> Block:
        rows = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            rows.append(" | ".join(cells))
        return Block(type="table", level=0, text="\n".join(rows))

    # ---------- Markdown ----------
    def _parse_markdown(self, path: Path) -> ParsedDocument:
        text = path.read_text(encoding="utf-8", errors="ignore")
        lines = text.splitlines()
        blocks: list[Block] = []
        title = path.stem
        buf: list[str] = []
        buf_type = "paragraph"

        def flush() -> None:
            nonlocal buf
            if buf:
                blocks.append(Block(type=buf_type, level=0, text="\n".join(buf)))
                buf = []

        for line in lines:
            m = re.match(r"^(#{1,6})\s+(.*)$", line)
            if m:
                flush()
                blocks.append(Block(type="heading", level=len(m.group(1)), text=m.group(2).strip()))
                continue
            if re.match(r"^\s*[-*+]\s+", line):
                flush()
                buf_type = "list_item"
                buf.append(re.sub(r"^\s*[-*+]\s+", "", line))
                continue
            if line.strip() == "":
                flush()
                buf_type = "paragraph"
                continue
            buf.append(line)
        flush()
        return ParsedDocument(title=title, blocks=blocks, raw_text=text, meta={"format": "markdown"})

    # ---------- HTML ----------
    def _parse_html(self, path: Path) -> ParsedDocument:
        from bs4 import BeautifulSoup  # 延迟导入

        html = path.read_text(encoding="utf-8", errors="ignore")
        soup = BeautifulSoup(html, "html.parser")
        title = soup.title.string if soup.title and soup.title.string else path.stem
        blocks: list[Block] = []
        for tag in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "table"]):
            name = tag.name
            text = tag.get_text(" ", strip=True)
            if not text:
                continue
            if name.startswith("h"):
                blocks.append(Block(type="heading", level=int(name[1]), text=text))
            elif name == "li":
                blocks.append(Block(type="list_item", level=1, text=text))
            elif name == "table":
                blocks.append(Block(type="table", level=0, text=text))
            else:
                blocks.append(Block(type="paragraph", level=0, text=text))
        return ParsedDocument(title=title, blocks=blocks, raw_text=soup.get_text(" ", strip=True), meta={"format": "html"})

    # ---------- TXT ----------
    def _parse_text(self, path: Path) -> ParsedDocument:
        text = path.read_text(encoding="utf-8", errors="ignore")
        blocks = self._blocks_from_pages([text])
        return ParsedDocument(title=path.stem, blocks=blocks, raw_text=text, page_count=1)

    # ---------- 辅助 ----------
    def _blocks_from_pages(self, pages: list[str]) -> list[Block]:
        blocks: list[Block] = []
        for page_no, text in enumerate(pages, start=1):
            for line in text.splitlines():
                line = line.strip()
                if not line:
                    continue
                m = HEADING_RE.match(line)
                if m and len(line) <= 60:
                    blocks.append(Block(type="heading", level=2, text=line, page=page_no))
                else:
                    blocks.append(Block(type="paragraph", level=0, text=line, page=page_no))
        return blocks
